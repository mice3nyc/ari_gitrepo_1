#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DMZ 영문 병기 docx 생성 (세션434, 6/5)
SPEC: _dev/dmz-layout/docs/SPEC.md §영문 병기 docx 생성

타깃 docx(한글 본문 92 ITEM, 워드 수작업 정리본)의 각 한글 ITEM 뒤에
EN CSV 영문 본문을 [English] 라벨과 함께 ITEM 블록 통째로 삽입한다.

배치(피터공 6/5): ITEM 블록 통째 / 구분 [English] 라벨 / 원본 덮어쓰기(사전 백업)
변환: 영문 HTML -> 텍스트 단락(<br>=단락분리, 빈칸 span->[X]) + <table>=docx 표

사용:
  python3 add_en_to_docx.py --verify-ko  # (참고) 변환 함수 한글 대조 — docx 수작업이라 100% 아님
  python3 add_en_to_docx.py --apply       # 백업 후 영문 삽입 + 저장
  python3 add_en_to_docx.py --check       # 삽입 결과 검증
  python3 add_en_to_docx.py --preview      # 삽입된 docx -> 한영 병기 txt 미리보기
"""
import csv, re, sys, os, shutil, unicodedata, glob
from copy import deepcopy
from html.parser import HTMLParser
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/p.air15/Neo-Obsi-Sync/Assets/incoming/통일부/본문 데이터 HTML"
DOCX_DIR = os.path.join(BASE, "영문번역/최종본문데이터영문포함")
EN_CSV = os.path.join(BASE, "사진링크용_본문채움_EN.csv")
KO_CSV = os.path.join(BASE, "사진링크용_본문채움_260529.csv")
BACKUP_DIR = os.path.join(DOCX_DIR, "_backup_pre_en")

TOPIC_FILES = {
    "DMZ 기본정보":   "260529_통일부_DMZ기본정보_영문포함.docx",
    "생태-환경":      "260529_통일부_생태_환경_영문포함.docx",
    "국가유산-문화재": "260529_통일부_국가유산_문화재_영문포함.docx",
    "DMZ의 사람들":   "260529_통일부_DMZ의 사람들_영문포함.docx",
    "평화 관광":      "260529_통일부_평화 관광_영문포함.docx",
}

HDR_RE = re.compile(r'^.+ / .+ / \d+\. ')
TITLE_RE = re.compile(r'^(.+?) / (.+) / \d+\. ')
EN_LABEL = "[English]"

def nfc(s): return unicodedata.normalize("NFC", s)

def norm(s):
    s = nfc(s)
    for ch in [" ", "\t", " ", "/", "·", "・", "|", ",", "-", "–", "—"]:
        s = s.replace(ch, "")
    return s

def squote(s):
    """비교용: 둥근/곧은 따옴표 통일"""
    return (s.replace("‘","'").replace("’","'")
             .replace("“",'"').replace("”",'"'))

# ---------- HTML 본문 -> 요소 리스트 [('para',txt) | ('table',rows)] ----------
class ElemParser(HTMLParser):
    BLOCK = {"div", "p", "li", "h1", "h2", "h3", "h4"}
    def __init__(self):
        super().__init__()
        self.elements = []
        self.cur = []
        self.in_blank = False; self.blank = None
        self.in_table = False; self.table = None; self.row = None; self.cell = None
    def _flush(self):
        t = "".join(self.cur).strip()
        if t: self.elements.append(("para", t))
        self.cur = []
    def handle_starttag(self, tag, attrs):
        a = dict(attrs)
        if tag == "table":
            self._flush(); self.in_table = True; self.table = []; return
        if self.in_table:
            if tag == "tr": self.row = []
            elif tag in ("td", "th"): self.cell = []
            elif tag == "br" and self.cell is not None: self.cell.append(" ")
            elif tag == "span" and a.get("class") == "dmz-blank":
                self.in_blank = True; self.blank = a.get("data-blank")
            return
        if tag == "br":
            self._flush(); return
        if tag == "span" and a.get("class") == "dmz-blank":
            self.in_blank = True; self.blank = a.get("data-blank"); return
        if tag in self.BLOCK:
            self._flush()
    def handle_endtag(self, tag):
        if self.in_table:
            if tag in ("td", "th") and self.cell is not None:
                self.row.append("".join(self.cell).strip()); self.cell = None
            elif tag == "tr" and self.row is not None:
                if any(c for c in self.row): self.table.append(self.row)
                self.row = None
            elif tag == "table":
                if self.table: self.elements.append(("table", self.table))
                self.in_table = False; self.table = None
            elif tag == "span" and self.in_blank:
                self.in_blank = False
            return
        if tag == "span" and self.in_blank:
            self.in_blank = False; return
        if tag in self.BLOCK:
            self._flush()
    def handle_data(self, d):
        if self.in_blank:
            tok = f"[{self.blank}]"
            (self.cell if (self.in_table and self.cell is not None) else self.cur).append(tok)
            return
        if self.in_table:
            if self.cell is not None: self.cell.append(d)
            return
        self.cur.append(d)
    def close(self):
        super().close(); self._flush()

def html_to_elements(html):
    p = ElemParser(); p.feed(html); p.close()
    out = []
    for kind, val in p.elements:
        if kind == "para":
            t = re.sub(r'[ \t ]+', ' ', nfc(val)).strip()
            out.append(("para", t))
        else:
            rows = [[re.sub(r'[ \t ]+',' ',nfc(c)).strip() for c in r] for r in val]
            out.append(("table", rows))
    return out

# ---------- CSV ----------
def load_csv(path):
    rows = list(csv.reader(open(path, encoding="utf-8")))
    items = {}
    for r in rows[2:]:
        if len(r) < 7 or not r[4].strip(): continue
        items[(norm(r[0]), norm(r[4]))] = {"주제": r[0], "타이틀": r[4], "html": r[6]}
    return items

# ---------- docx ITEM 분할 (body element 순서) ----------
def split_items(doc):
    body = doc.element.body
    P, SECT = qn('w:p'), qn('w:sectPr')
    p_map = {p._p: p for p in doc.paragraphs}
    items, cur = [], None
    for el in list(body):
        if el.tag == SECT:
            continue
        if el.tag == P:
            para = p_map.get(el)
            txt = para.text.strip() if para else ""
            if HDR_RE.match(txt):
                m = TITLE_RE.match(txt)
                cur = {"title": m.group(2).strip() if m else "", "header_el": el,
                       "ko_paras": [], "els": [el], "end_el": None}
                items.append(cur)
            elif cur is not None:
                if txt: cur["ko_paras"].append(txt)
                cur["els"].append(el)
        elif cur is not None:           # table 등
            cur["els"].append(el)
    for i, it in enumerate(items):
        it["end_el"] = items[i+1]["header_el"] if i+1 < len(items) else None
    return items

def strip_page_breaks(els):
    """ITEM 영역 단락들에서 명시적 페이지 브레이크 run(<w:br type='page'>) 제거"""
    n = 0
    for el in els:
        for br in list(el.iter(qn('w:br'))):
            if br.get(qn('w:type')) == 'page':
                br.getparent().remove(br); n += 1
    return n

def _mk_pagebreak_el(doc):
    p = doc.add_paragraph()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
    p.add_run()._r.append(br)
    el = p._p; el.getparent().remove(el)
    return el

# ---------- 삽입 헬퍼 ----------
# 표는 별도 임시 문서에서 생성(타깃 docx의 소수 twips 마진 버그 우회) 후 element 복제
_TMP = docx.Document()

def _mk_para_el(doc, text, bold=False):
    p = doc.add_paragraph()              # body 끝(sectPr 앞)에 생성
    p.style = doc.styles["normal"]
    run = p.add_run(text)
    if bold: run.bold = True
    el = p._p
    el.getparent().remove(el)            # 떼어내 호출부에서 위치 결정
    return el

def _mk_table_el(rows):
    cols = max(len(r) for r in rows)
    t = _TMP.add_table(rows=len(rows), cols=cols)
    try: t.style = "Table Grid"
    except Exception: pass
    for ri, r in enumerate(rows):
        for ci in range(cols):
            t.cell(ri, ci).text = r[ci] if ci < len(r) else ""
    el = deepcopy(t._tbl)
    t._tbl.getparent().remove(t._tbl)    # 임시 문서 정리
    return el

def insert_en_block(doc, anchor_el, elements, add_pagebreak=True):
    """anchor_el(다음 ITEM 헤더 또는 sectPr) 앞에 [English]+영문 요소 삽입.
    add_pagebreak면 영문 블록 끝에 페이지 브레이크 단락 추가 (ITEM 끝에서 페이지 나눔)."""
    nodes = [_mk_para_el(doc, ""),           # 한글과 구분 빈 줄
             _mk_para_el(doc, EN_LABEL)]     # [English]
    for idx, (kind, val) in enumerate(elements):
        if idx == 0:
            # 첫 요소 = 타이틀: [English] 바로 아래(빈 줄 없이), 볼드
            nodes.append(_mk_table_el(val) if kind == "table"
                         else _mk_para_el(doc, val, bold=True))
        else:
            nodes.append(_mk_para_el(doc, ""))   # 요소 사이 빈 줄
            nodes.append(_mk_table_el(val) if kind == "table" else _mk_para_el(doc, val))
    for node in nodes:
        anchor_el.addprevious(node)          # 항상 anchor 앞으로 (순서 유지)
    if add_pagebreak:
        anchor_el.addprevious(_mk_pagebreak_el(doc))

def do_apply():
    os.makedirs(BACKUP_DIR, exist_ok=True)
    en = load_csv(EN_CSV)
    total_items, total_en, missing = 0, 0, []
    for topic, fn in TOPIC_FILES.items():
        src = os.path.join(DOCX_DIR, fn)
        bak = os.path.join(BACKUP_DIR, fn)
        if not os.path.exists(bak):
            shutil.copy2(src, bak)         # 원본 백업(최초 1회)
        doc = docx.Document(bak)           # 항상 깨끗한 백업에서 시작(재실행 안전)
        tail = doc.element.body.find(qn('w:sectPr'))  # 마지막 ITEM 삽입 기준
        items = split_items(doc)
        for it in items:
            total_items += 1
            key = (norm(topic), norm(it["title"]))
            if key not in en:
                missing.append((topic, it["title"])); continue
            strip_page_breaks(it["els"])   # ITEM 한글 영역의 기존 페이지 브레이크 제거
            els = html_to_elements(en[key]["html"])
            is_last = it["end_el"] is None
            anchor = tail if is_last else it["end_el"]
            # 영문 뒤 페이지 브레이크: 마지막 ITEM은 문서 끝이라 생략
            insert_en_block(doc, anchor, els, add_pagebreak=not is_last)
            total_en += 1
        doc.save(src)
        print(f"  저장: {fn} (ITEM {len(items)})")
    print(f"[apply] ITEM {total_items} / 영문삽입 {total_en} / 매칭실패 {len(missing)}")
    for m in missing: print("   매칭실패:", m)

def do_check():
    en = load_csv(EN_CSV)
    print(f"{'파일':32} ITEM  EN라벨  표  빈칸[X]  백스페이스  비NFC")
    g_items=g_en=0
    for topic, fn in TOPIC_FILES.items():
        path = os.path.join(DOCX_DIR, fn)
        doc = docx.Document(path)
        items = split_items(doc)
        full = "\n".join(p.text for p in doc.paragraphs)
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells: full += "\n"+c.text
        en_labels = sum(1 for p in doc.paragraphs if p.text.strip()==EN_LABEL)
        ntable = len(doc.tables)
        nblank = len(re.findall(r'\[[A-Z]\]', full))
        nbs = full.count("\x08")
        non_nfc = sum(1 for p in doc.paragraphs if p.text != nfc(p.text))
        print(f"{fn:32} {len(items):4}  {en_labels:5}  {ntable:2}  {nblank:6}  {nbs:9}  {non_nfc:5}")
        g_items+=len(items); g_en+=en_labels
    print(f"{'합계':32} {g_items:4}  {g_en:5}   (ITEM=EN라벨=92 이어야 정상)")

def do_preview():
    """삽입 결과를 한영 병기 txt로 추출(피터공 육안용)"""
    out_dir = os.path.join(DOCX_DIR, "_preview_txt")
    os.makedirs(out_dir, exist_ok=True)
    import subprocess
    for fn in TOPIC_FILES.values():
        src = os.path.join(DOCX_DIR, fn)
        dst = os.path.join(out_dir, fn.replace(".docx",".txt"))
        subprocess.run(["textutil","-convert","txt","-output",dst,src], check=True)
        print("  preview:", dst)

def do_verify_ko():
    """참고용: 변환 함수가 한글 docx와 얼마나 맞는지(수작업 정리라 100% 불가)"""
    ko = load_csv(KO_CSV)
    total, exact = 0, 0
    for topic, fn in TOPIC_FILES.items():
        # 깨끗한 백업이 있으면 거기서(이미 영문 삽입됐을 수 있음)
        path = os.path.join(BACKUP_DIR, fn)
        if not os.path.exists(path): path = os.path.join(DOCX_DIR, fn)
        doc = docx.Document(path)
        for it in split_items(doc):
            total += 1
            key = (norm(topic), norm(it["title"]))
            if key not in ko: continue
            conv = [squote(t) for k,t in html_to_elements(ko[key]["html"]) if k=="para"]
            docx_p = [squote(re.sub(r'[ \t ]+',' ',nfc(t)).strip()) for t in it["ko_paras"]]
            if conv == docx_p: exact += 1
    print(f"[verify-ko 참고] 텍스트 단락 완전일치 {exact}/{total} "
          f"(docx는 워드 수작업 정리본이라 라벨·표로 차이 — 영문 삽입과 무관)")

def main():
    mode = sys.argv[1] if len(sys.argv) > 1 else "--check"
    {"--apply": do_apply, "--check": do_check, "--preview": do_preview,
     "--verify-ko": do_verify_ko}.get(mode, lambda: print("모드:", mode))()

if __name__ == "__main__":
    main()
